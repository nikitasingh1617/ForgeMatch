import os
import json
import re
from typing import List, Dict, Any, Optional
import PyPDF2
import docx
import io
from datetime import datetime

def parse_resume_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Main entry point - detects file type and parses accordingly
    
    Args:
        file_content: Raw bytes of the uploaded file
        filename: Original filename with extension
    
    Returns:
        Dict containing candidate profile data
    
    Raises:
        ValueError: If file format is unsupported or parsing fails
    """
    # 1. Get file extension (e.g., '.pdf', '.docx')
    file_extension = os.path.splitext(filename)[1].lower()
    
    # 2. Route to appropriate parser
    if file_extension == '.json':
        return parse_json(file_content)
    elif file_extension == '.jsonl':
        return parse_jsonl(file_content)
    elif file_extension == '.pdf':
        return parse_pdf(file_content)
    elif file_extension in ['.docx', '.doc']:
        return parse_docx(file_content)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Please use JSON, JSONL, PDF, or DOCX.")
    
def parse_json(file_content: bytes) -> Dict[str, Any]:
    """Parse JSON file - can be single candidate or array"""
    try:
        data = json.loads(file_content.decode('utf-8'))
        
        # If it's a list, return first candidate (or you could handle multiple)
        # Our ranking engine expects a single candidate or list
        if isinstance(data, list):
            return data[0] if data else {}
        return data
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON format: {str(e)}")

def parse_jsonl(file_content: bytes) -> Dict[str, Any]:
    """Parse JSONL (JSON Lines) file - each line is a JSON object"""
    try:
        # For JSONL, we'll take the first line as the candidate
        # (In production, you might want to handle multiple)
        lines = file_content.decode('utf-8').strip().split('\n')
        if not lines:
            raise ValueError("Empty file")
        
        # Filter out empty lines
        lines = [line for line in lines if line.strip()]
        
        # Parse the first line
        return json.loads(lines[0])
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSONL format: {str(e)}")
    
def parse_pdf(file_content: bytes) -> Dict[str, Any]:
    """Extract text from PDF and convert to candidate profile"""
    try:
        # Create a file-like object from bytes
        pdf_file = io.BytesIO(file_content)
        
        # Create PDF reader
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            try:
                pdf_reader.decrypt('')
            except:
                raise ValueError("PDF is encrypted and cannot be decrypted.")
        
        # Extract text from all pages
        text = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                print(f"Warning: Could not extract text from page {page_num}: {str(e)}")
                continue
        
        if not text.strip():
            raise ValueError("No text could be extracted from PDF. The PDF might be scanned or image-based.")
        
        # Convert extracted text to candidate profile
        return extract_candidate_from_text(text)
        
    except PyPDF2.errors.PdfReadError as e:
        raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")    
    
def parse_docx(file_content: bytes) -> Dict[str, Any]:
    """Extract text from DOCX and convert to candidate profile"""
    try:
        # Create a file-like object from bytes
        docx_file = io.BytesIO(file_content)
        
        # Create DOCX document
        doc = docx.Document(docx_file)
        
        # Extract text from all paragraphs
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text:
                        text += para.text + "\n"
            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text:
                        text += para.text + "\n"
        
        if not text.strip():
            raise ValueError("No text could be extracted from DOCX.")
        
        # Convert extracted text to candidate profile
        return extract_candidate_from_text(text)
        
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")    

def extract_candidate_from_text(text: str) -> Dict[str, Any]:
    """
    Extract candidate information from plain text resume.
    This uses pattern matching and section detection.
    
    Args:
        text: Raw text from resume
    
    Returns:
        Structured candidate profile dictionary
    """
    # Initialize candidate structure matching your ranking.py format
    candidate = {
        "candidate_id": f"resume_{abs(hash(text)) % 100000:05d}",
        "profile": {
            "headline": "",
            "summary": "",
            "current_title": "",
            "current_industry": "",
            "years_of_experience": 0,
            "location": "",
            "country": ""
        },
        "skills": [],
        "career_history": [],
        "education": [],
        "projects": [],
        "certifications": []
    }
    
    # Split into lines and clean
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Section detection keywords
    sections = {
        'skills': ['skills', 'technologies', 'tech stack', 'programming languages', 
                   'expertise', 'core competencies', 'technical skills', 'languages'],
        'experience': ['experience', 'work history', 'employment', 'career', 
                       'work experience', 'professional experience', 'employment history'],
        'education': ['education', 'qualifications', 'academic', 'degree', 
                      'university', 'college', 'school'],
        'projects': ['projects', 'portfolio', 'personal projects', 'side projects'],
        'summary': ['summary', 'profile', 'about me', 'professional summary', 'objective'],
        'certifications': ['certifications', 'certificates', 'accreditations']
    }
    
    current_section = 'header'
    skill_set = set()  # To avoid duplicates
    
    # First pass: detect if we have structured sections
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Detect section headers
        for section, keywords in sections.items():
            if any(kw in line_lower for kw in keywords):
                # This line is a section header
                current_section = section
                break
    
    # Second pass: process content by section
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        line_lower = line.lower()
        
        # Skip section headers
        is_section_header = False
        for keywords in sections.values():
            if any(kw in line_lower for kw in keywords):
                is_section_header = True
                break
        if is_section_header:
            continue
        
        # Process based on current section
        if current_section == 'header':
            # Try to detect name and title
            if not candidate['profile']['headline']:
                # First non-empty line is often the name or headline
                candidate['profile']['headline'] = line
            else:
                # Second line might be current title
                if not candidate['profile']['current_title'] and len(line.split()) <= 6:
                    candidate['profile']['current_title'] = line
        
        elif current_section == 'summary':
            if candidate['profile']['summary']:
                candidate['profile']['summary'] += " " + line
            else:
                candidate['profile']['summary'] = line
        
        elif current_section == 'skills':
            # Extract skills - split by commas, bullets, or spaces
            # Also handle skills separated by "|" or "/"
            skills = re.split(r'[,•·●◆▪▫\n|/]+', line)
            for skill in skills:
                skill = skill.strip()
                if skill and len(skill) > 1 and skill not in skill_set:  # Avoid single characters and duplicates
                    skill_set.add(skill)
                    candidate['skills'].append({
                        "name": skill,
                        "proficiency": "unknown"  # Default if not specified
                    })
        
        elif current_section == 'experience':
            # Try to parse job entries
            # Pattern: "Title at Company" or "Company - Title"
            if ' at ' in line:
                parts = line.split(' at ', 1)
                title = parts[0].strip()
                company = parts[1].strip()
                candidate['career_history'].append({
                    "title": title,
                    "company": company,
                    "description": "",
                    "industry": "",
                    "start_date": "",
                    "end_date": ""
                })
            elif ' - ' in line and any(word in line_lower for word in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant', 'architect']):
                # Alternative format: "Company - Title"
                parts = line.split(' - ', 1)
                company = parts[0].strip()
                title = parts[1].strip()
                candidate['career_history'].append({
                    "title": title,
                    "company": company,
                    "description": "",
                    "industry": "",
                    "start_date": "",
                    "end_date": ""
                })
            elif '|' in line and any(word in line_lower for word in ['engineer', 'developer', 'manager', 'analyst', 'designer']):
                # Format: "Company | Title"
                parts = line.split('|', 1)
                company = parts[0].strip()
                title = parts[1].strip()
                candidate['career_history'].append({
                    "title": title,
                    "company": company,
                    "description": "",
                    "industry": "",
                    "start_date": "",
                    "end_date": ""
                })
            else:
                # If we have a current job, append description to it
                if candidate['career_history']:
                    if candidate['career_history'][-1]['description']:
                        candidate['career_history'][-1]['description'] += " " + line
                    else:
                        candidate['career_history'][-1]['description'] = line
        
        elif current_section == 'education':
            # Simple education parsing - try to extract degree and institution
            if ' in ' in line or ' at ' in line or ',' in line:
                # Try to parse "Degree in Field at University" or "Degree, University"
                degree = line
                institution = ""
                if ' at ' in line:
                    parts = line.split(' at ', 1)
                    degree = parts[0].strip()
                    institution = parts[1].strip()
                elif ',' in line:
                    parts = line.split(',', 1)
                    degree = parts[0].strip()
                    institution = parts[1].strip()
                candidate['education'].append({
                    "degree": degree,
                    "institution": institution,
                    "field_of_study": ""
                })
            else:
                candidate['education'].append({
                    "degree": line,
                    "institution": "",
                    "field_of_study": ""
                })
        
        elif current_section == 'projects':
            # Try to parse project info
            if ':' in line:
                name, desc = line.split(':', 1)
                candidate['projects'].append({
                    "name": name.strip(),
                    "description": desc.strip(),
                    "tech_stack": []
                })
            elif ' - ' in line:
                # Format: "Project Name - Description"
                parts = line.split(' - ', 1)
                candidate['projects'].append({
                    "name": parts[0].strip(),
                    "description": parts[1].strip() if len(parts) > 1 else "",
                    "tech_stack": []
                })
            else:
                # Add as description to last project if exists
                if candidate['projects']:
                    if candidate['projects'][-1]['description']:
                        candidate['projects'][-1]['description'] += " " + line
                    else:
                        candidate['projects'][-1]['description'] = line
                else:
                    # Create a new project with the line as description
                    candidate['projects'].append({
                        "name": f"Project {len(candidate['projects']) + 1}",
                        "description": line,
                        "tech_stack": []
                    })
        
        elif current_section == 'certifications':
            # Try to extract certification name and issuer
            if ' from ' in line or ' by ' in line:
                if ' from ' in line:
                    parts = line.split(' from ', 1)
                    name = parts[0].strip()
                    issuer = parts[1].strip()
                else:
                    parts = line.split(' by ', 1)
                    name = parts[0].strip()
                    issuer = parts[1].strip()
                candidate['certifications'].append({
                    "name": name,
                    "issuer": issuer
                })
            else:
                candidate['certifications'].append({
                    "name": line,
                    "issuer": ""
                })
    
    # Extract years of experience using regex
    year_patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)',
        r'(\d+)\s*\+\s*(?:years?|yrs?)',
        r'(\d+)-(\d+)\s*(?:years?|yrs?)',
        r'(\d+)\s*(?:years?|yrs?)\s+experience',
        r'(?:experience|exp)\s*[:\-]?\s*(\d+)\s*(?:years?|yrs?)'
    ]
    
    for pattern in year_patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            if isinstance(matches[0], tuple):
                # Range like "3-5 years"
                candidate['profile']['years_of_experience'] = int(matches[0][1] or matches[0][0])
            else:
                candidate['profile']['years_of_experience'] = int(matches[0])
            break
    
    # Extract location (common patterns)
    location_patterns = [
        r'(?:location|based in|from)\s*[:\-]?\s*([A-Za-z\s,]+?)(?:\n|$)',
        r'([A-Za-z\s,]+?)\s*,\s*[A-Z]{2}\s*(?:\n|$)',  # City, State
        r'(?:city|state|country)\s*[:\-]?\s*([A-Za-z\s,]+?)(?:\n|$)'
    ]
    
    for pattern in location_patterns:
        location_match = re.search(pattern, text, re.IGNORECASE)
        if location_match:
            candidate['profile']['location'] = location_match.group(1).strip().title()
            break
    
    # Extract email (common pattern)
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        candidate['profile']['email'] = email_match.group(0)
    
    # Extract phone number (common patterns)
    phone_patterns = [
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # (123) 456-7890 or 123-456-7890
        r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'  # +1 234 567 8900
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            candidate['profile']['phone'] = phone_match.group(0)
            break
    
    # Clean up empty entries
    candidate['career_history'] = [j for j in candidate['career_history'] if j.get('title') or j.get('company')]
    candidate['education'] = [e for e in candidate['education'] if e.get('degree')]
    candidate['projects'] = [p for p in candidate['projects'] if p.get('name') or p.get('description')]
    candidate['certifications'] = [c for c in candidate['certifications'] if c.get('name')]
    
    # Ensure skills are unique
    seen_skills = set()
    unique_skills = []
    for skill in candidate['skills']:
        if skill['name'] not in seen_skills:
            seen_skills.add(skill['name'])
            unique_skills.append(skill)
    candidate['skills'] = unique_skills
    
    return candidate

# ─── Helper function to parse multiple candidates from JSONL ──
def parse_jsonl_multiple(file_content: bytes) -> List[Dict[str, Any]]:
    """
    Parse JSONL file and return all candidates as a list
    Useful for bulk uploads
    """
    try:
        lines = file_content.decode('utf-8').strip().split('\n')
        if not lines:
            raise ValueError("Empty file")
        
        # Filter out empty lines
        lines = [line for line in lines if line.strip()]
        
        candidates = []
        for line in lines:
            try:
                candidate = json.loads(line)
                candidates.append(candidate)
            except json.JSONDecodeError as e:
                print(f"Warning: Skipping invalid JSON line: {str(e)}")
                continue
        
        return candidates
    except Exception as e:
        raise ValueError(f"Failed to parse JSONL: {str(e)}")

# ─── Helper function to parse multiple candidates from JSON array ──
def parse_json_multiple(file_content: bytes) -> List[Dict[str, Any]]:
    """
    Parse JSON file containing an array of candidates
    """
    try:
        data = json.loads(file_content.decode('utf-8'))
        
        if isinstance(data, list):
            return data
        else:
            # Single candidate - wrap in list
            return [data]
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON format: {str(e)}")